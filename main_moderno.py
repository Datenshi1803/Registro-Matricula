# =========================================================
# SISTEMA ACADÉMICO - REGISTRO Y ANÁLISIS DE DATOS  V1.1
# =========================================================

# -------------------------
# IMPORTACIONES
# -------------------------
import os
import datetime
import sqlite3

import tkinter as tk
from tkinter import messagebox

import ttkbootstrap as ttk
from ttkbootstrap.constants import *

import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

from docx import Document
from modals import Modal
from PIL import Image, ImageTk, ImageDraw  
#Locales
from modals import Modal
from modales.crud_niveles import abrir_crud_niveles
from modales.crud_profesores import abrir_crud_profesores
from modales.crud_asignaturas import abrir_crud_asignaturas
from modales.crud_asignaciones import abrir_asignacion_profesor_asignaturas
from modales.crud_relacion_nivel_profesor import abrir_asignacion_niveles_profesor



# =========================================================
# CONFIGURACIÓN GLOBAL / CONSTANTES
# =========================================================

# Tamaño de ventana
ANCHO_VENTANA = 1100
ALTO_VENTANA = 750

# Ruta de la imagen de fondo
IMAGEN_PATH = "fondo_cefine.png"

# Nivel de opacidad para el fondo (0 = invisible, 255 = opaco)
OPACIDAD_NIVEL = 90


# =========================================================
# VENTANA PRINCIPAL (ttkbootstrap)
# =========================================================

# Usamos ttk.Window para aplicar un tema (flatly)
ventana = ttk.Window(title="Sistema Académico - Registro", themename="flatly")

# Centrar ventana en la pantalla
ancho_pantalla = ventana.winfo_screenwidth()
alto_pantalla = ventana.winfo_screenheight()
x = (ancho_pantalla // 2) - (ANCHO_VENTANA // 2)
y = (alto_pantalla // 2) - (ALTO_VENTANA // 2)
ventana.geometry(f"{ANCHO_VENTANA}x{ALTO_VENTANA}+{x}+{y}")


# =========================================================
# FUNCIONES DE UTILIDAD (FONDO, VALIDACIÓN, EXPORTACIÓN)
# =========================================================

def establecer_fondo(contenedor, imagen_path):
    """Coloca una imagen de fondo con opacidad en un contenedor."""
    try:
        img = Image.open(imagen_path).convert("RGBA")

        contenedor.update_idletasks()
        current_width = contenedor.winfo_width()
        current_height = contenedor.winfo_height()

        img_redim = img.resize((current_width, current_height))

        # Aplicar opacidad
        alpha = img_redim.split()[3]
        alpha = Image.eval(alpha, lambda x: x * (OPACIDAD_NIVEL / 255))
        img_redim.putalpha(alpha)

        fondo_img = ImageTk.PhotoImage(img_redim)

        fondo_label = tk.Label(contenedor, image=fondo_img, borderwidth=0)
        fondo_label.image = fondo_img
        fondo_label.place(x=0, y=0, relwidth=1, relheight=1)
        fondo_label.lower()

        def resize_image(event):
            try:
                new_width = event.width
                new_height = event.height

                img_resized = Image.open(imagen_path).convert("RGBA").resize((new_width, new_height))

                alpha_resized = img_resized.split()[3]
                alpha_resized = Image.eval(alpha_resized, lambda x: x * (OPACIDAD_NIVEL / 255))
                img_resized.putalpha(alpha_resized)

                new_image = ImageTk.PhotoImage(img_resized)

                fondo_label.config(image=new_image)
                fondo_label.image = new_image
                fondo_label.lower()
            except Exception:
                pass

        contenedor.bind('<Configure>', resize_image)

    except FileNotFoundError:
        print(f"Advertencia: Archivo de fondo no encontrado en {imagen_path}. No se aplicará fondo.")
    except Exception as e:
        print(f"Error al cargar la imagen de fondo con opacidad: {e}")


def solo_numeros(valor):
    """Valida que solo se ingresen números (o vacío para permitir borrar)."""
    if valor == "" or valor.isdigit():
        return True
    else:
        ventana.bell()
        return False


def exportar_a_excel(datos, tipo):
    """Exporta los datos de la gráfica a un archivo Excel."""
    nombre_archivo = f"grafica_{tipo.replace(' ', '_')}.xlsx"
    df_export = pd.DataFrame(datos)
    df_export.to_excel(nombre_archivo)
    messagebox.showinfo("Exportado", f"Gráfica exportada como {nombre_archivo}")


def exportar_a_word(datos, tipo):
    """Exporta los datos de la gráfica a un archivo Word."""
    nombre_archivo = f"grafica_{tipo.replace(' ', '_')}.docx"
    doc = Document()
    doc.add_heading("Reporte de Gráfica", level=1)
    doc.add_paragraph(f"Tipo: {tipo}")

    tabla = doc.add_table(rows=1, cols=2)
    hdr_cells = tabla.rows[0].cells
    hdr_cells[0].text = "Categoría"
    hdr_cells[1].text = "Valor"

    for index, value in datos.items():
        row_cells = tabla.add_row().cells
        row_cells[0].text = str(index)
        row_cells[1].text = str(value)

    doc.save(nombre_archivo)
    messagebox.showinfo("Exportado", f"Gráfica exportada como {nombre_archivo}")
    
def cargar_profesores_por_nivel(event):
    nivel = widget_map["Nivel"].get()

    conn = sqlite3.connect("registro.db")
    cursor = conn.cursor()
    cursor.execute("""
        SELECT profesores.nombre
        FROM profesores
        JOIN nivel_profesor ON profesores.id = nivel_profesor.id_profesor
        JOIN niveles ON niveles.id = nivel_profesor.id_nivel
        WHERE niveles.nombre = ?
        ORDER BY profesores.nombre
    """, (nivel,))
    profesores = [p[0] for p in cursor.fetchall()]
    conn.close()

    widget_map["Profesor"]["values"] = profesores

    # Limpiar asignaturas
    widget_map["Asignatura"]["values"] = []
    widget_map["Asignatura"].set("")

def cargar_asignaturas_por_profesor(event):
    profesor = widget_map["Profesor"].get()

    conn = sqlite3.connect("registro.db")
    cursor = conn.cursor()
    cursor.execute("""
        SELECT asignaturas.nombre
        FROM asignaturas
        JOIN profesor_asignatura ON asignaturas.id = profesor_asignatura.id_asignatura
        JOIN profesores ON profesores.id = profesor_asignatura.id_profesor
        WHERE profesores.nombre = ?
        ORDER BY asignaturas.nombre
    """, (profesor,))
    asignaturas = [a[0] for a in cursor.fetchall()]
    conn.close()

    widget_map["Asignatura"]["values"] = asignaturas

    if len(asignaturas) == 1:
        widget_map["Asignatura"].set(asignaturas[0])


# =========================================================
# VARIABLES GLOBALES PARA EXPORTACIÓN DE GRÁFICAS
# =========================================================

ultimo_dato = pd.Series(dtype=float)
ultimo_titulo = ""


# =========================================================
# FUNCIONES PRINCIPALES (DB, GRÁFICAS, FORMULARIO)
# =========================================================

def generar_grafica(tipo):
    """Genera la gráfica según el tipo seleccionado y los filtros aplicados."""
    conn = sqlite3.connect("registro.db")
    df = pd.read_sql_query("SELECT * FROM registros", conn)
    conn.close()

    if df.empty:
        messagebox.showwarning("Sin datos", "No hay registros disponibles.")
        return
    
    # Aplicar filtros
    nivel = nivel_filtro.get().strip()
    asignatura = asignatura_filtro.get().strip()
    anio = anio_filtro.get().strip()

    if nivel:
        df = df[df['nivel'].str.contains(nivel, case=False)]
    if asignatura:
        df = df[df['asignatura'].str.contains(asignatura, case=False)]
    if anio:
        df = df[df['anio'].astype(str) == anio]

    if df.empty:
        messagebox.showwarning("Sin resultados", "No se encontraron registros con los filtros seleccionados.")
        return

    fig, ax = plt.subplots(figsize=(6, 4.5))

    # Lógica según tipo de gráfica
    if tipo == "Aprobados por nivel y año":
        datos = df.groupby("nivel")["aprobados"].sum()
        titulo = "Aprobados por Nivel y Año"
    elif tipo == "Reprobados por nivel y año":
        datos = df.groupby("nivel")["reprobados"].sum()
        titulo = "Reprobados por Nivel y Año"
    elif tipo == "Aprobados por nivel y asignatura":
        datos = df.groupby("asignatura")["aprobados"].sum()
        titulo = "Aprobados por Asignatura"
    elif tipo == "Reprobados por nivel y asignatura":
        datos = df.groupby("asignatura")["reprobados"].sum()
        titulo = "Reprobados por Asignatura"
    else:
        messagebox.showinfo("Selecciona un tipo", "Por favor selecciona un tipo de gráfica válido.")
        return

    if datos.empty or datos.sum() == 0:
        messagebox.showinfo("Sin datos", "No hay valores para graficar en esta categoría.")
        return

    total = datos.sum()

    def formato_etiqueta(porcentaje, valores=datos):
        valor = int(round(porcentaje * total / 100.0))
        return f"{valor} ({porcentaje:.1f}%)"

    wedges, texts, autotexts = ax.pie(
        datos,
        labels=datos.index,
        autopct=lambda p: formato_etiqueta(p),
        startangle=90,
        textprops={'color': 'black', 'fontsize': 9}
    )

    ax.set_title(titulo, fontsize=14, fontweight='bold')
    ax.axis('equal')

    ax.legend(
        loc="center left",
        bbox_to_anchor=(1, 0.5),
        title="Categorías",
        labels=[f"{cat}: {val}" for cat, val in zip(datos.index, datos)],
        fontsize=9
    )

    fig.tight_layout(pad=2.0)
    fig.subplots_adjust(right=0.75, top=0.9, bottom=0.1)

    # Limpiar contenido previo
    for widget in frame_preview.winfo_children():
        widget.destroy()

    canvas = FigureCanvasTkAgg(fig, master=frame_preview)
    canvas.draw()
    canvas_widget = canvas.get_tk_widget()
    canvas_widget.pack(expand=True, fill="both", padx=20, pady=15)

    def on_resize(event):
        fig.tight_layout()
        canvas.draw()

    frame_preview.bind("<Configure>", on_resize)

    # Actualizar globales para exportación
    global ultimo_dato, ultimo_titulo
    ultimo_dato, ultimo_titulo = datos, titulo

    # Mostrar botones de exportación
    boton_excel.grid()
    boton_word.grid()

def crear_tablas_relaciones():
        conn = sqlite3.connect("registro.db")
        cursor = conn.cursor()

        # Tabla niveles 
        cursor.execute("""
        CREATE TABLE IF NOT EXISTS niveles (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT UNIQUE NOT NULL
        )
        """)

        # Profesores
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS profesores (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                nombre TEXT NOT NULL
            )
        """)

        # Asignaturas
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS asignaturas (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                nombre TEXT UNIQUE NOT NULL
            )
        """)

        # Relación profesor - asignatura
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS profesor_asignatura (
                id_profesor INTEGER,
                id_asignatura INTEGER,
                FOREIGN KEY(id_profesor) REFERENCES profesores(id),
                FOREIGN KEY(id_asignatura) REFERENCES asignaturas(id)
            )
        """)

        # Relación nivel - profesor (el nivel tiene varios profesores)
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS nivel_profesor (
                id_nivel INTEGER,
                id_profesor INTEGER,
                FOREIGN KEY(id_nivel) REFERENCES niveles(id),
                FOREIGN KEY(id_profesor) REFERENCES profesores(id)
            )
        """)

        conn.commit()
        conn.close()

crear_tablas_relaciones()

def guardar_datos():
    """Guarda los datos del formulario en la base de datos SQLite."""
    try:
        # Validación básica
        if not all(v.get() for _, v in campos[:3]):
            messagebox.showwarning("Advertencia", "Los campos Nivel, Profesor y Asignatura no pueden estar vacíos.")
            return

        conn = sqlite3.connect("registro.db")
        cursor = conn.cursor()
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS registros (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                nivel TEXT, profesor TEXT, asignatura TEXT, anio INTEGER,
                trimestre INTEGER, turno TEXT, hombres INTEGER, mujeres INTEGER,
                aprobados INTEGER, reprobados INTEGER, reprobados_fecha INTEGER,
                sin_asistencia INTEGER, sin_calificacion INTEGER, sin_profesor INTEGER,
                retirados INTEGER
            )
        """)

        valores = [v.get() if v.get() else None for _, v in campos]

        cursor.execute("""
            INSERT INTO registros (
                nivel, profesor, asignatura, anio, trimestre, turno,
                hombres, mujeres, aprobados, reprobados, reprobados_fecha,
                sin_asistencia, sin_calificacion, sin_profesor, retirados
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, valores)

        conn.commit()
        conn.close()

        messagebox.showinfo("Éxito", "Datos guardados correctamente en la base de datos.")
        limpiar_campos()

    except Exception as e:
        messagebox.showerror("Error", f"No se pudo guardar el registro:\n{e}")


def limpiar_campos():
    """Limpia todos los campos del formulario."""
    for _, var in campos:
        var.set("")


# =========================================================
# PESTAÑAS PRINCIPALES
# =========================================================

tabs = ttk.Notebook(ventana)
tab_registro = ttk.Frame(tabs, padding=15)
tab_graficas = ttk.Frame(tabs, padding=15)

tabs.add(tab_registro, text="📋 Registro")
tabs.add(tab_graficas, text="📊 Gráficas")
tabs.pack(expand=1, fill="both", padx=10, pady=10)


# =========================================================
# FRAMES PRINCIPALES (REGISTRO / GRÁFICAS)
# =========================================================

# Frame Registro
frame = ttk.Frame(tab_registro, padding=20, borderwidth=1, relief="flat", bootstyle='light')
frame.pack(expand=True, fill="both")

# Frame Gráficas
frame_graf = ttk.Frame(tab_graficas, padding=20, bootstyle='light')
frame_graf.pack(expand=True, fill="both")

# Fondo solo en gráficas por ahora
establecer_fondo(frame_graf, IMAGEN_PATH)
# Si quieres fondo también en registro, descomenta:
# establecer_fondo(frame, IMAGEN_PATH)

# Responsividad
frame.columnconfigure(1, weight=1)
frame.columnconfigure(3, weight=1)
frame.rowconfigure(18, weight=1)

frame_graf.columnconfigure(1, weight=1)
frame_graf.columnconfigure(3, weight=1)
frame_graf.rowconfigure(4, weight=5)


# =========================================================
# SECCIÓN GRÁFICAS: FILTROS, BOTONES, PREVIEW
# =========================================================

# Variables para filtros
nivel_filtro = tk.StringVar()
asignatura_filtro = tk.StringVar()
anio_filtro = tk.StringVar()
tipo_filtro = tk.StringVar()

# Título
ttk.Label(
    frame_graf,
    text="Análisis de Resultados",
    font=("Segoe UI", 14, "bold"),
    style='primary.TLabel'
).grid(row=0, column=0, columnspan=4, pady=10)

# Filtros
ttk.Label(frame_graf, text="Nivel:").grid(row=1, column=0, sticky="e", padx=5, pady=5)
ttk.Entry(frame_graf, textvariable=nivel_filtro, width=25).grid(row=1, column=1, padx=5, pady=5, sticky="ew")

ttk.Label(frame_graf, text="Asignatura:").grid(row=1, column=2, sticky="e", padx=5, pady=5)
ttk.Entry(frame_graf, textvariable=asignatura_filtro, width=25).grid(row=1, column=3, padx=5, pady=5, sticky="ew")

ttk.Label(frame_graf, text="Año:").grid(row=2, column=0, sticky="e", padx=5, pady=5)
ttk.Entry(frame_graf, textvariable=anio_filtro, width=25).grid(row=2, column=1, padx=5, pady=5, sticky="ew")

ttk.Label(frame_graf, text="Tipo de gráfica:").grid(row=2, column=2, sticky="e", padx=5, pady=5)
ttk.Combobox(
    frame_graf,
    textvariable=tipo_filtro,
    values=[
        "Aprobados por nivel y año",
        "Reprobados por nivel y año",
        "Aprobados por nivel y asignatura",
        "Reprobados por nivel y asignatura"
    ],
    state="readonly",
    width=23
).grid(row=2, column=3, padx=5, pady=5, sticky="ew")

# Botón generar
boton_generar = ttk.Button(
    frame_graf,
    text="📈 Generar gráfica",
    bootstyle=PRIMARY,
    command=lambda: generar_grafica(tipo_filtro.get())
)
boton_generar.grid(row=3, column=0, pady=15, padx=5, sticky="w")

# Botones de exportación (se muestran luego de generar)
boton_excel = ttk.Button(
    frame_graf,
    text="📤 Exportar a Excel",
    bootstyle=INFO,
    command=lambda: exportar_a_excel(ultimo_dato, ultimo_titulo)
)
boton_word = ttk.Button(
    frame_graf,
    text="📝 Exportar a Word",
    bootstyle=INFO,
    command=lambda: exportar_a_word(ultimo_dato, ultimo_titulo)
)

boton_excel.grid(row=3, column=2, pady=15, padx=5, sticky="e")
boton_word.grid(row=3, column=3, pady=15, padx=5, sticky="e")

boton_excel.grid_remove()
boton_word.grid_remove()

# Área de vista previa de gráficas
frame_preview = ttk.Frame(frame_graf)
frame_preview.grid(row=4, column=0, columnspan=4, pady=10, sticky="nsew")


# =========================================================
# SECCIÓN REGISTRO: HEADER (LOGO + TÍTULO)
# =========================================================

frame_header = ttk.Frame(frame)
frame_header.grid(row=0, column=0, columnspan=4, pady=10, sticky="ew")
ttk.Button(
    frame_header,
    text="Administrar profesores, niveles y asignaturas",
    bootstyle=INFO,
    command=lambda: abrir_modal_catalogos()
).pack(side="right", padx=10)

# Logo
try:
    logo_path = "logo.png"
    if os.path.exists(logo_path):
        logo_img = Image.open(logo_path).convert("RGBA")
        logo_img.thumbnail((80, 80), Image.Resampling.LANCZOS)
        logo_tk = ImageTk.PhotoImage(logo_img)

        logo_label = tk.Label(frame_header, image=logo_tk, borderwidth=0, bg='white')
        logo_label.image = logo_tk
        logo_label.pack(side="left", padx=10)
except Exception as e:
    print(f"No se pudo cargar el logo: {e}")

# Título
ttk.Label(
    frame_header,
    text="Registro Académico",
    font=("Segoe UI", 16, "bold"),
    style='primary.TLabel'
).pack(side="left", padx=10, expand=True)


# =========================================================
# SECCIÓN REGISTRO: CAMPOS Y VALIDACIONES
# =========================================================

widget_map = {}


# Variables
campos = [
    ("Nivel", tk.StringVar()), ("Profesor", tk.StringVar()),
    ("Asignatura", tk.StringVar()), ("Año", tk.StringVar()),
    ("Trimestre", tk.StringVar()), ("Turno", tk.StringVar()),
    ("Hombres", tk.StringVar()), ("Mujeres", tk.StringVar()),
    ("Aprobados", tk.StringVar()), ("Reprobados", tk.StringVar()),
    ("Reprobados a la fecha", tk.StringVar()), ("Sin asistencia", tk.StringVar()),
    ("Sin calificación", tk.StringVar()), ("Sin profesor", tk.StringVar()),
    ("Retirados", tk.StringVar())
]

col_izq = campos[:6]   
col_der = campos[6:]   

# Subtítulos
ttk.Label(
    frame,
    text="Datos Generales",
    font=("Segoe UI", 11, "bold"),
    style='info.TLabel'
).grid(row=1, column=0, sticky="w", pady=5)

ttk.Label(
    frame,
    text="Estadísticas",
    font=("Segoe UI", 11, "bold"),
    style='info.TLabel'
).grid(row=1, column=2, sticky="w", pady=5)

# Validación numérica
vcmd = ventana.register(solo_numeros)

# Campos izquierda y derecha
for i, (label, var) in enumerate(col_izq, start=2):
    ttk.Label(frame, text=label).grid(row=i, column=0, sticky="e", padx=5, pady=5)

    # ---- NIVEL ----
    if label == "Nivel":
        conn = sqlite3.connect("registro.db")
        cursor = conn.cursor()
        cursor.execute("SELECT nombre FROM niveles ORDER BY nombre")
        niveles = [n[0] for n in cursor.fetchall()]
        conn.close()

        cb = ttk.Combobox(
            frame, textvariable=var,
            values=niveles, state="readonly",
            width=23, bootstyle=PRIMARY
        )
        cb.grid(row=i, column=1, sticky="ew", padx=5, pady=5)
        cb.bind("<<ComboboxSelected>>", cargar_profesores_por_nivel)

        widget_map[label] = cb
        continue

    # ---- PROFESOR ----
    if label == "Profesor":
        cb = ttk.Combobox(
            frame, textvariable=var,
            values=[], state="readonly",
            width=23, bootstyle=PRIMARY
        )
        cb.grid(row=i, column=1, sticky="ew", padx=5, pady=5)
        cb.bind("<<ComboboxSelected>>", cargar_asignaturas_por_profesor)

        widget_map[label] = cb
        continue

    # ---- ASIGNATURA ----
    if label == "Asignatura":
        cb = ttk.Combobox(
            frame, textvariable=var,
            values=[], state="readonly",
            width=23, bootstyle=PRIMARY
        )
        cb.grid(row=i, column=1, sticky="ew", padx=5, pady=5)
        widget_map[label] = cb
        continue

    # ---- TURNO ----
    if label == "Turno":
        cb = ttk.Combobox(
            frame, textvariable=var,
            values=["Mañana", "Tarde"], state="readonly",
            width=23, bootstyle=PRIMARY
        )
        cb.grid(row=i, column=1, sticky="ew", padx=5, pady=5)
        widget_map[label] = cb
        continue

    # ---- TRIMESTRE ----
    if label == "Trimestre":
        cb = ttk.Combobox(
            frame, textvariable=var,
            values=["1", "2", "3"], state="readonly",
            width=23, bootstyle=PRIMARY
        )
        cb.grid(row=i, column=1, sticky="ew", padx=5, pady=5)
        widget_map[label] = cb
        continue

    # ---- AÑO ----
    if label == "Año":
        anios = [str(a) for a in range(2020, datetime.date.today().year + 2)]
        cb = ttk.Combobox(
            frame, textvariable=var,
            values=anios, state="readonly",
            width=23, bootstyle=PRIMARY
        )
        cb.grid(row=i, column=1, sticky="ew", padx=5, pady=5)
        cb.set(str(datetime.date.today().year))

        widget_map[label] = cb
        continue

    # =========================================================
    # CAMPOS DE LA COLUMNA DERECHA
    # =========================================================

for i, (label, var) in enumerate(col_der, start=2):
    ttk.Label(frame, text=label).grid(row=i, column=2, sticky="e", padx=5, pady=5)

    # Campos numéricos
    if label in ["Hombres", "Mujeres", "Aprobados", "Reprobados",
                 "Reprobados a la fecha", "Sin asistencia",
                 "Sin calificación", "Sin profesor", "Retirados"]:
        entry = ttk.Entry(
            frame, textvariable=var, width=25,
            validate="key", validatecommand=(vcmd, "%P")
        )
        entry.grid(row=i, column=3, sticky="ew", padx=5, pady=5)
        widget_map[label] = entry
        continue

    # Campos generales
    entry = ttk.Entry(frame, textvariable=var, width=25)
    entry.grid(row=i, column=3, sticky="ew", padx=5, pady=5)
    widget_map[label] = entry


# Separador
ttk.Separator(frame, orient="horizontal", bootstyle=SECONDARY).grid(
    row=19, column=0, columnspan=4, sticky="ew", pady=10
)


# =========================================================
# SECCIÓN REGISTRO: BOTONES
# =========================================================

ttk.Button(
    frame,
    text="Guardar",
    width=18,
    bootstyle=PRIMARY,
    command=guardar_datos
).grid(row=20, column=0, columnspan=2, pady=10)

ttk.Button(
    frame,
    text="Limpiar",
    width=18,
    bootstyle=DANGER,
    command=limpiar_campos
).grid(row=20, column=2, columnspan=2, pady=10)


# =========================================================
# FUNCIONES PARA CRUD ADICIONALES (NIVELES, PROFESORES, ASIGNATURAS)
# =========================================================


def abrir_modal_catalogos():
    modal = Modal(
        parent=ventana,
        title="Administración de Catálogos",
        width=400,
        height=350
    )

    frame = modal.body

    ttk.Button(
        frame,
        text="Administrar Niveles",
        bootstyle="primary",
        width=30,
        command=lambda: (modal.destroy(), abrir_crud_niveles(ventana))
    ).pack(pady=10)

    ttk.Button(
        frame,
        text="Administrar Profesores",
        bootstyle="info",
        width=30,
        command=lambda: (modal.destroy(), abrir_crud_profesores(ventana))
    ).pack(pady=10)

    ttk.Button(
        frame,
        text="Administrar Asignaturas",
        bootstyle="secondary",
        width=30,
        command=lambda: (modal.destroy(), abrir_crud_asignaturas(ventana))
    ).pack(pady=10)

    ttk.Button(
        frame,
        text="Asignar Asignaturas a Profesores",
        bootstyle="warning",
        width=30,
        command=lambda: (modal.destroy(), abrir_asignacion_profesor_asignaturas(ventana))
    ).pack(pady=10)
    
    ttk.Button(
    frame,
    text="Asignar Niveles a Profesores",
    bootstyle=WARNING,
    width=30,
    command=lambda: (modal.destroy(), abrir_asignacion_niveles_profesor(ventana))
    ).pack(pady=10)



# Botón para abrir el menú de catálogos
ttk.Button(ventana, text="Abrir Catálogos", command=abrir_modal_catalogos).pack(pady=20)

# =========================================================
# INICIO DE LA APLICACIÓN
# =========================================================

ventana.mainloop()
